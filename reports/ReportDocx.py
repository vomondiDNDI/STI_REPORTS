#import docx
#import os #to run/open document automatically


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT
from docx.shared import Inches,Pt
import psycopg2
#from SQLFileQuery import *
#from SQLFileCRF import *
#from LabSummarySQL import *
#from AELabelling import *
from docx.oxml.shared import OxmlElement, qn
from docxtpl import DocxTemplate
import datetime
from Settings import PATHS,DB
#from expected_crf import *
import os, shutil
import math
import pandas as pd


#################################################################

document = Document()

p1 = document.add_paragraph('')
run =p1.add_run('STI - Zoliflodacin phase III')
run.italic = False
run.bold = False
font = run.font
font.name = 'Calibri'
font.size = Pt(14)

paragraph_format = p1.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format.space_before = Pt(48)
paragraph_format.space_after = Pt(18)
p2 = document.add_paragraph('')
run =p2.add_run('PROGRESS REPORT')
run.italic = False
run.bold = True
font = run.font
font.name = 'Calibri'
font.size = Pt(16)

paragraph_format = p2.paragraph_format
paragraph_format.space_before = Pt(48)
paragraph_format.space_after = Pt(18)
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3 = document.add_paragraph('By the 5th of every month')
paragraph_format = p3.paragraph_format
paragraph_format.space_before = Pt(48)
paragraph_format.space_after = Pt(4)
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
now = datetime.datetime.now()
p4 = document.add_paragraph()
run =p4.add_run('(With data received as of '+ now.strftime("%d-%B-%Y")+')')
run.italic = True
run.bold = False
font = run.font
font.name = 'Calibri'
font.size = Pt(10)
paragraph_format = p4.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_page_break()
p1 = document.add_paragraph('Table of contents')


paragraph = document.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
#instrText.text = 'TOC \o "1-3" \h \z \u'  # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p


document.add_page_break()
new_section = document.add_section()
new_section.orientation = WD_ORIENT.LANDSCAPE
document.save('monthly_report1.docx')
#document.save(PATHS['TEMP_FOLDER']+'monthly_report1.docx')
#os.system('monthly_report1.docx')

doc1 = Document('monthly_report1.docx')
#document.save(PATHS['TEMP_FOLDER']+'monthly_report1.docx')
#doc2 = Document('schedule_template1.docx')
doc2 = Document('ScheduleEvents_template.docx')
sec = doc2.add_section()
sec.orientation = WD_ORIENT.LANDSCAPE
for element in doc2.element.body:
    doc1 .element.body.append(element)
#doc1.add_page_break()

doc1.save('monthly_report2.docx')

document = Document()
"""Start of Randomizer Tables"""
document.add_heading('Table 2: Enrollment by Site', level=1)
"""
"""
# create an instance of a word document
#document = docx.Document()
records = [
    ['Mbagathi Health Centre-01', '50'],
    ['Special Treatment Centre Casino Health Centre-01', '70'],
    ['Coast Provincial General Hospital-01', '100'],
    ['International Centre for Reproductive Health-01', '120'],
    ['Homa-Bay County Referral Hospital-01', '175']
    ['Total Enrolled', '475']
]

table_main = document.add_table(rows=1, cols=2)
table_main.allow_autofit = True
table_main.style='Table Grid'
hdr_cells = table_main.rows[0].cells #first row
#hdr_cells[0].text = 'SiteCode' #first row text/heading
hdr_cells[0].text = 'SiteName'
hdr_cells[1].text = 'Enrolled'


for SiteName, Enrolled in records:
    row_cells = table_main.add_row().cells
    #row_cells[0].text = str(SiteCode)
    row_cells[0].text = SiteName
    row_cells[1].text = Enrolled

#################################################################
#document.save('monthly_report3.docx')
#os.system('monthly_report3.docx')
document.add_page_break()
document.add_heading('Table 3: Query Management Status', level=1)
# create an instance of a word document
#document = docx.Document()
records = [
    ['Mbagathi Health Centre', '1050','980','20','8'],
    ['Special Treatment Centre Casino Health Centre', '1050','980','30','8'],
    ['Coast Provincial General Hospital', '1050','980','15','8'],
    ['International Centre for Reproductive Health', '1050','980','10','8'],
    ['Homa-Bay County Referral Hospital', '1050','980','50','8'],
    ['Total', '3150','300','24','24']
]

table_main = document.add_table(rows=1, cols=5, style="Light Grid Accent 6")
table_main.allow_autofit = True
table_main.style='Table Grid'
hdr_cells = table_main.rows[0].cells #first row
hdr_cells[0].text = 'Site Name'
hdr_cells[1].text = 'All Queries'
hdr_cells[2].text = 'All Open Queries'
hdr_cells[3].text = 'Unresolved Queries'
hdr_cells[4].text = 'Closed Queries'


for SiteName, AllQueries, AllOpenQueries,UnresolvedQueries, ClosedQueries in records:
    row_cells = table_main.add_row().cells
    row_cells[0].text = SiteName
    row_cells[1].text = AllQueries
    row_cells[2].text = AllOpenQueries
    row_cells[3].text = UnresolvedQueries
    row_cells[4].text = ClosedQueries

#################################################################
#document.save('monthly_report3.docx')
#os.system('monthly_report3.docx')
document.add_page_break()
# create an instance of a word document
#document = docx.Document()
records = [
    ['1016', 'fever of unknown origin','2018-11-24','2018-11-24','Fatal','None'],
    ['1021', 'multiorgan failure','2019-02-12','2019-02-12','Fatal','None'],
    ['1041', 'acute kidney injury','2019-12-08','2019-12-08','Resolved','None']
]
document.add_heading('Table 4: Serious Adverse Events Listings', level=1)

table_main = document.add_table(rows=1, cols=6, style="Light Grid Accent 6")
table_main.allow_autofit = True
table_main.style='Table Grid'
hdr_cells = table_main.rows[0].cells #first row
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'SAE Description'
hdr_cells[2].text = 'Start Date'
hdr_cells[3].text = 'End Date'
hdr_cells[4].text = 'Intensity'
hdr_cells[5].text = 'Outcome'


for ID, SAEDescription, StartDate, EndDate, Intensity, Outcome in records:
    row_cells = table_main.add_row().cells
    row_cells[0].text = str(ID)
    row_cells[1].text = SAEDescription
    row_cells[2].text = StartDate
    row_cells[3].text = EndDate
    row_cells[4].text = Intensity
    row_cells[5].text = Outcome

#################################################################


document.add_page_break()
# create an instance of a word document
#document = docx.Document()
records = [
    ['716', 'Nausea','2018-11-24','2018-11-24','Resolved','Yes','None','Yes'],
    ['700', 'Vomiting','2019-02-12','2019-02-12','Resolved','Yes','None','Yes'],
    ['724', 'Diarrhoea','2019-12-08','2019-12-08','Resolved','Yes','None','Yes']
]
document.add_heading('Table 5: Adverse Events Listings', level=1)

table_main = document.add_table(rows=1, cols=8, style="Light Grid Accent 6")
table_main.allow_autofit = True
table_main.style='Table Grid'
hdr_cells = table_main.rows[0].cells #first row
hdr_cells[0].text = 'AE Number'
hdr_cells[1].text = 'Adverse Event'
hdr_cells[2].text = 'Start Date'
hdr_cells[3].text = 'Intensity'
hdr_cells[4].text = 'Relation to Participant'
hdr_cells[5].text = 'Outcome'
hdr_cells[6].text = 'End Date'
hdr_cells[7].text = 'Was Concomitant Medication Given?'


for AEno, AE , StartDate, Intensity,RParticipant, Outcome, EndDate ,ConcomitantGiven in records:
    row_cells = table_main.add_row().cells
    row_cells[0].text = str(AEno)
    row_cells[1].text = AE
    row_cells[2].text = StartDate
    row_cells[3].text = Intensity
    row_cells[4].text = RParticipant
    row_cells[5].text = Outcome
    row_cells[6].text = EndDate
    row_cells[7].text = ConcomitantGiven

document.save('monthly_report3.docx')

doc3 = Document('monthly_report2.docx')
doc4 = Document('monthly_report3.docx')
for element in doc4.element.body:
    doc3 .element.body.append(element)
doc3.save('STImonthly_report.docx')
os.system('STImonthly_report.docx')


